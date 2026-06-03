from docx import Document
from docx.shared import Pt


DOCX_PATH = "CSB22044_46_ENDTERM_8_REPORT_app_extended.docx"


def add_before(anchor, text="", bold=False, italic=False):
    paragraph = anchor.insert_paragraph_before(text)

    if paragraph.runs:
        run = paragraph.runs[0]
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)

    return paragraph


def add_heading(anchor, text):
    paragraph = anchor.insert_paragraph_before()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return paragraph


def add_bullet(anchor, text):
    paragraph = anchor.insert_paragraph_before(f"- {text}")
    paragraph.paragraph_format.left_indent = Pt(18)
    return paragraph


def main():
    document = Document(DOCX_PATH)

    chapter_6_matches = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip().upper().startswith("CHAPTER 6: RESULTS"):
            chapter_6_matches.append(paragraph)

    if not chapter_6_matches:
        raise RuntimeError("Could not locate CHAPTER 6 anchor.")

    toc_chapter_6 = chapter_6_matches[0]
    chapter_6 = chapter_6_matches[-1]

    toc_entries = [
        "5.5 Application Implementation: HyperClusterAI Dashboard",
        "5.6 Application Workflow",
        "5.7 Clustering Methods Implemented in the Application",
        "5.8 CNN Autoencoder Integration",
        "5.9 Dashboard Modules",
        "5.10 Evaluation and Export System",
        "5.11 Performance Optimization",
        "5.12 Application Significance",
    ]

    for entry in toc_entries:
        add_before(toc_chapter_6, entry)

    add_heading(chapter_6, "5.5 Application Implementation: HyperClusterAI Dashboard")
    add_before(
        chapter_6,
        "In addition to the experimental notebook-based implementation, an interactive application named HyperClusterAI was developed using Streamlit. The purpose of the application is to convert the clustering workflow into a practical analysis platform where RGB and hyperspectral images can be uploaded, processed, clustered, evaluated, visualized, compared, and exported through a single interface.",
    )
    add_before(
        chapter_6,
        "The dashboard extends the research implementation by making the system usable for demonstrations, comparative experimentation, and visual interpretation. Instead of running every method manually through separate scripts, the user can select a clustering method, configure parameters, execute the pipeline, and inspect the output through organized dashboard tabs.",
    )

    add_heading(chapter_6, "5.6 Application Workflow")
    add_before(
        chapter_6,
        "The application follows a complete spatial-spectral segmentation workflow:",
    )
    for item in [
        "Upload RGB or hyperspectral image data in JPG, PNG, NPY, or MAT format.",
        "Optionally upload a ground truth mask in NPY or MAT format.",
        "Automatically detect whether the input is RGB or hyperspectral based on the number of bands.",
        "Configure the number of clusters, patch size, PCA components, latent dimension, single band index, and false-color band indices.",
        "Select either a single clustering method or multiple methods for comparison.",
        "Run the selected clustering pipeline.",
        "Visualize the original image, feature representation, cluster map, metrics, spectral signatures, t-SNE feature space, exports, and runtime performance.",
    ]:
        add_bullet(chapter_6, item)

    add_heading(chapter_6, "5.7 Clustering Methods Implemented in the Application")
    add_before(
        chapter_6,
        "The current application supports the following methods:",
    )
    for item in [
        "Single Band KMeans: Uses one selected spectral band as a grayscale baseline.",
        "False Color KMeans: Uses three selected spectral bands as a false-color feature representation.",
        "KMeans: Performs clustering on standardized full pixel vectors.",
        "PCA + KMeans: Reduces spectral dimensionality using PCA before applying KMeans.",
        "Spatial-Spectral Clustering: Extracts local patches around each pixel, applies scaling and PCA, and clusters the resulting spatial-spectral features.",
        "Dense Autoencoder Clustering: Uses flattened spatial-spectral patch features, PCA reduction, dense autoencoder training, latent feature extraction, and KMeans clustering.",
        "CNN Autoencoder Clustering: Reduces spectral bands using PCA, extracts patch cubes, preserves local patch structure using convolutional layers, learns latent features, and applies KMeans clustering.",
        "Deep Embedded Clustering (DEC): Pretrains an autoencoder, initializes clusters using KMeans, and refines cluster assignments through KL-divergence-based target distribution learning.",
    ]:
        add_bullet(chapter_6, item)

    add_heading(chapter_6, "5.8 CNN Autoencoder Integration")
    add_before(
        chapter_6,
        "The application was extended with a CNN Autoencoder method to better align with spatial-spectral hyperspectral clustering. Unlike the dense autoencoder, which flattens patch information into a one-dimensional vector, the CNN Autoencoder preserves the two-dimensional neighborhood structure of each local patch. This allows the model to learn local texture, edge continuity, neighborhood consistency, and spatial-spectral patterns more naturally.",
    )
    add_before(
        chapter_6,
        "The CNN Autoencoder pipeline used in the application is:",
    )
    add_before(
        chapter_6,
        "Input Image -> Pixel Scaling -> PCA Band Reduction -> Patch Cube Extraction -> CNN Autoencoder -> Latent Embedding -> KMeans -> Cluster Map",
        italic=True,
    )
    add_before(
        chapter_6,
        "The encoder uses convolutional layers followed by a dense latent layer. The decoder reconstructs the patch cube from the learned latent representation. After training, the encoder output is used as the feature representation for clustering. This method is computationally heavier than the dense autoencoder but is more suitable for spatial-spectral feature learning.",
    )

    add_heading(chapter_6, "5.9 Dashboard Modules")
    add_before(
        chapter_6,
        "The Streamlit dashboard is organized into the following modules:",
    )
    for item in [
        "Original Image: Displays the uploaded RGB image or a selected RGB-style preview of hyperspectral bands.",
        "Feature View: Displays reduced or learned feature representations when available.",
        "Cluster Results: Shows the generated segmentation map using a spectral color map.",
        "Metrics: Shows Silhouette Score and, when ground truth is available, Accuracy, Cohen's Kappa, NMI, confusion matrix, and class-wise accuracy.",
        "Comparison: Runs multiple selected methods and compares their quantitative performance.",
        "Spectral Signatures: Plots the mean spectral signature of each predicted cluster to verify whether clusters correspond to meaningful material groups.",
        "Latent Space: Uses t-SNE to visualize feature separability in two dimensions.",
        "Exports: Provides downloadable cluster maps, metric tables, comparison graphs, confusion matrices, class accuracy graphs, and report summaries.",
        "Performance: Displays runtime, memory usage, image size, processing mode, method complexity, dataset type, and CPU/GPU backend availability.",
    ]:
        add_bullet(chapter_6, item)

    add_heading(chapter_6, "5.10 Evaluation and Export System")
    add_before(
        chapter_6,
        "The application supports both internal and external clustering evaluation. When ground truth labels are not available, the Silhouette Score is used as an internal validation metric. When ground truth is provided, external metrics such as Accuracy, Cohen's Kappa, and Normalized Mutual Information are computed. Since cluster labels are arbitrary in unsupervised learning, Hungarian matching is used to align predicted cluster labels with ground truth labels before computing Accuracy and Kappa.",
    )
    add_before(
        chapter_6,
        "The export system improves practical usability by allowing users to download cluster maps, metric CSV files, confusion matrix images, class-wise accuracy graphs, comparison reports, and method summary reports. This makes the application useful for documentation, presentation, and further analysis.",
    )

    add_heading(chapter_6, "5.11 Performance Optimization")
    add_before(
        chapter_6,
        "To improve Streamlit loading time, TensorFlow-based modules are loaded lazily. This means the dashboard opens without immediately importing TensorFlow. TensorFlow is imported only when deep learning methods such as Dense Autoencoder, CNN Autoencoder, or DEC are executed. Classical methods such as KMeans, PCA + KMeans, Single Band KMeans, False Color KMeans, and Spatial-Spectral clustering therefore load faster and are suitable for quick demonstrations.",
    )
    add_before(
        chapter_6,
        "For faster demonstrations, smaller patch sizes such as 3x3 or 5x5, PCA components between 10 and 30, and moderate latent dimensions are recommended. CNN Autoencoder and DEC should preferably be used on smaller images or reduced hyperspectral datasets because they require more computation and memory.",
    )

    add_heading(chapter_6, "5.12 Application Significance")
    add_before(
        chapter_6,
        "The HyperClusterAI application strengthens the project by combining research experimentation with practical usability. It enables interactive method selection, parameter tuning, visualization, evaluation, comparison, and export. As a result, the project is not limited to offline experiments but becomes a usable hyperspectral image clustering platform.",
    )

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
